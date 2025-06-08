# Step 3: Assemble Word Document with python-docx
from docx import Document

def assemble_word_report(summary):
    doc = Document("templates/report_template.docx")

    # Replace placeholders
    for p in doc.paragraphs:
        if "{MONTH}" in p.text:
            p.text = p.text.replace("{MONTH}", "May 2025")

    # Fill summary table
    table = doc.tables[0]
    for region, amount in summary.itertuples(index=False):
        row_cells = table.add_row().cells
        row_cells[0].text = str(region)
        row_cells[1].text = f"${amount:,.2f}"

    doc_output = "output/report_2025-05.docx"
    doc.save(doc_output)
    print("Word report created:", doc_output)
